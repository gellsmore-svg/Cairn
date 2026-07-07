"""Export integration hooks for downstream docx/PDF/HTML tooling.

Built-in support (install 'cairn-lang[export]' for docx/pdf):
- "html": self-contained HTML page (no extra deps)
- "docx": Microsoft Word document (via python-docx)
- "pdf": PDF document (via fpdf2)

Example:
    from cairn.render import render_plan, export_view
    view = render_plan("...", output_format="markdown")
    pdf_bytes = export_view(view, "pdf", options={"title": "My Plan"})
"""

from __future__ import annotations

from collections.abc import Callable
from html import escape
from io import BytesIO
from typing import Any

from cairn.render.model import RenderResult

Exporter = Callable[[RenderResult, dict[str, Any]], bytes]

_EXPORTERS: dict[str, Exporter] = {}


def _html_exporter(result: RenderResult, options: dict[str, Any]) -> bytes:
    """Simple self-contained HTML exporter."""
    title = options.get("title", "Cairn View")
    body = escape(result.body).replace("\n", "<br>\n")
    if result.footnotes:
        notes = "<h3>Notes</h3><ul>" + "".join(f"<li>{escape(n)}</li>" for n in result.footnotes) + "</ul>"
    else:
        notes = ""
    html = f"""<!DOCTYPE html>
<html lang="{result.language}">
<head>
<meta charset="utf-8">
<title>{escape(title)}</title>
<style>
body {{ font-family: system-ui, sans-serif; line-height: 1.5; max-width: 800px; margin: 2em auto; padding: 0 1em; }}
h1, h2, h3 {{ color: #222; }}
pre, code {{ background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; }}
</style>
</head>
<body>
<h1>{escape(title)}</h1>
<p><em>Profile: {escape(result.profile)} | Language: {escape(result.language)}</em></p>
<div>{body}</div>
{notes}
</body>
</html>"""
    return html.encode("utf-8")


def _docx_exporter(result: RenderResult, options: dict[str, Any]) -> bytes:
    """Export to .docx using python-docx (optional dep)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for docx export: pip install 'cairn-lang[export]'"
        ) from exc

    doc = Document()
    title = options.get("title", f"Cairn View - {result.profile}")
    heading = doc.add_heading(title, level=0)
    para = doc.add_paragraph()
    run = para.add_run(f"Profile: {result.profile} | Language: {result.language}")
    run.italic = True

    for line in result.body.splitlines():
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(6)

    if result.footnotes:
        doc.add_heading("Notes", level=1)
        for n in result.footnotes:
            doc.add_paragraph(n, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_exporter(result: RenderResult, options: dict[str, Any]) -> bytes:
    """Export to PDF using fpdf2 (optional dep)."""
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise ImportError(
            "fpdf2 is required for pdf export: pip install 'cairn-lang[export]'"
        ) from exc

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=16)
    title = options.get("title", f"Cairn View - {result.profile}")
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 8, f"Profile: {result.profile} | Language: {result.language}", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", size=11)
    for line in result.body.splitlines():
        # Simple wrapping
        pdf.multi_cell(0, 6, line)
        pdf.ln(1)

    if result.footnotes:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Notes", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", size=10)
        for n in result.footnotes:
            pdf.multi_cell(0, 6, f"- {n}")
            pdf.ln(1)

    return pdf.output()


def register_exporter(format_name: str, exporter: Exporter) -> None:
    """Register a binary exporter (e.g. docx, pdf, html) for ``export_view``."""
    _EXPORTERS[format_name.lower()] = exporter


def registered_exporters() -> list[str]:
    return sorted(_EXPORTERS)


def export_view(result: RenderResult, format_name: str, options: dict[str, Any] | None = None) -> bytes:
    """Export a rendered view via a registered plugin exporter."""
    key = format_name.lower()
    if key not in _EXPORTERS:
        known = ", ".join(sorted(_EXPORTERS)) or "(none registered)"
        raise NotImplementedError(
            f"No exporter registered for {format_name!r}. "
            f"Register one with cairn.render.export.register_exporter (built-ins: {known}). "
            f"Or use cairn-render with markdown/text/json/mermaid."
        )
    return _EXPORTERS[key](result, options or {})


# Register built-ins after defs
register_exporter("html", _html_exporter)
register_exporter("docx", _docx_exporter)
register_exporter("pdf", _pdf_exporter)